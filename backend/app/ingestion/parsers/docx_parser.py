"""Parse DOCX files into plain text using python-docx.

Extracts text from paragraphs and tables, preserving document structure:
- Headings are detected by style name and prefixed with markdown-style markers
  (#, ##, ###, ...) so the chunker treats them as natural section boundaries.
- Table rows are extracted as pipe-delimited text.
- Paragraph-level text is kept in reading order.
"""

import re
from pathlib import Path


class DocxParser:
    """Extract structured plain text from .docx (Office Open XML) files.

    Preserves document hierarchy by detecting Word heading styles
    (Heading 1 through Heading 9) and emitting markdown-style heading
    prefixes. This ensures the RecursiveChunker (which splits first on
    ``\\n\\n``) respects section boundaries.
    """

    SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".docx"})
    SUPPORTED_TYPES: frozenset[str] = frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )

    # Regex to match paragraph styles like "Heading 1", "Heading 2", … "Heading 9"
    _HEADING_STYLE_RE: re.Pattern[str] = re.compile(r"^[Hh]eading\s+(\d{1,2})$")

    @staticmethod
    def supports(filename: str) -> bool:
        """Check if this parser supports the given filename."""
        ext = Path(filename).suffix.lower()
        return ext in DocxParser.SUPPORTED_EXTENSIONS

    @staticmethod
    async def parse(content: bytes, filename: str) -> str:
        """Extract structured text from DOCX bytes.

        Args:
            content: Raw DOCX file bytes.
            filename: Original filename (used only for error messages).

        Returns:
            Plain text with markdown-style heading prefixes and
            paragraph/table separators, suitable for recursive chunking.

        Raises:
            ValueError: If the file is not a valid DOCX or has no text.
        """
        ext = Path(filename).suffix.lower()
        if ext not in DocxParser.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file extension '{ext}'. "
                f"Supported: {', '.join(sorted(DocxParser.SUPPORTED_EXTENSIONS))}"
            )

        if not content:
            raise ValueError("Empty DOCX: nothing to extract.")

        from io import BytesIO

        import docx

        try:
            doc = docx.Document(BytesIO(content))
        except Exception as exc:
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc

        parts: list[str] = []

        # Extract paragraphs, detecting heading levels
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            level = DocxParser._heading_level(para)
            if level is not None:
                # Emit heading as "# " + text (or "## ", "### ", etc.)
                prefix = "#" * level
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

        # Extract table content as structured rows
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    table_rows.append(" | ".join(row_texts))
            if table_rows:
                # Blank line before table block so the chunker sees a boundary
                parts.append("")
                parts.extend(table_rows)
                parts.append("")

        result = "\n\n".join(parts).strip()
        if not result:
            raise ValueError("DOCX contains no extractable text.")

        return result

    @staticmethod
    def _heading_level(para) -> int | None:
        """Return the heading level (1–9) if the paragraph is a heading,
        or ``None`` otherwise.

        Inspects ``para.style.name`` against the heading regex.
        """
        style_name = getattr(para.style, "name", "")
        if not style_name:
            return None
        m = DocxParser._HEADING_STYLE_RE.match(style_name)
        if m:
            level = int(m.group(1))
            return min(level, 9)  # clamp to 9 for sanity
        return None
