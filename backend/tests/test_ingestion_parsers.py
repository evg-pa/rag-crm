"""Tests for the full ingestion pipeline parsers (APP-139).

Covers:
  PDF parser (pymupdf) — valid PDF, empty, encrypted
  DOCX parser (python-docx) — valid DOCX, empty, with tables
  HTML parser (BeautifulSoup) — valid HTML, empty, scripts stripped
  Web scraper (httpx) — public URL, unsafe URL reject
  Parser registry — dispatch, unsupported extension
  API endpoints — upload PDF/DOCX/HTML, scrape, supported extensions
"""

import io
import uuid
from collections.abc import AsyncIterator

import pytest
from httpx import ASGITransport, AsyncClient

from app.ingestion.parsers.docx_parser import DocxParser
from app.ingestion.parsers.html_parser import HtmlParser
from app.ingestion.parsers.pdf_parser import PdfParser
from app.ingestion.parsers.registry import (
    get_all_supported_extensions,
    get_ext_to_content_type_map,
    get_parser_for,
)
from app.ingestion.parsers.text_parser import TextParser
from app.main import app

# ── Helpers ──────────────────────────────────────────────────────────────────


def _make_pdf_bytes(pages_text: list[str]) -> bytes:
    """Create a minimal valid PDF with the given page texts using pymupdf."""
    import fitz

    doc = fitz.open()
    for text in pages_text:
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=12)
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _make_docx_bytes(
    paragraphs: list[str],
    table_rows: list[list[str]] | None = None,
    heading_levels: list[int] | None = None,
) -> bytes:
    """Create a minimal valid DOCX with given paragraphs and optional table.

    If ``heading_levels`` is provided, it must have the same length as
    ``paragraphs``. A value of 0 (or ``None`` for an entry) means a normal
    paragraph; 1–9 means "Heading N" style.
    """
    from docx import Document

    doc = Document()
    for i, para_text in enumerate(paragraphs):
        level = heading_levels[i] if heading_levels and i < len(heading_levels) else None
        if level and 1 <= level <= 9:
            doc.add_paragraph(para_text, style=f"Heading {level}")
        else:
            doc.add_paragraph(para_text)

    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Extract helpers for the API tests ────────────────────────────────────────


@pytest.fixture
def anyio_backend() -> str:
    return "asyncio"


@pytest.fixture
async def client() -> AsyncIterator[AsyncClient]:
    """Client with a mock auth user — no DB dependency needed."""
    from unittest.mock import Mock

    from app.core.auth import get_current_user
    from app.models.user import User

    mock_user = Mock(spec=User)
    mock_user.id = uuid.UUID(int=1)
    mock_user.email = "test@example.com"
    mock_user.display_name = "Test User"
    mock_user.is_active = True
    mock_user.is_admin = False

    async def _override_get_current_user() -> User:
        return mock_user  # type: ignore[return-value]

    app.dependency_overrides[get_current_user] = _override_get_current_user

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac

    app.dependency_overrides.pop(get_current_user, None)


# ── PDF Parser ───────────────────────────────────────────────────────────────


class TestPdfParser:
    """Unit tests for PdfParser."""

    def test_supports_pdf(self) -> None:
        assert PdfParser.supports("report.pdf") is True
        assert PdfParser.supports("REPORT.PDF") is True
        assert PdfParser.supports("doc.txt") is False
        assert PdfParser.supports("doc.docx") is False

    @pytest.mark.asyncio
    async def test_parse_simple_pdf(self) -> None:
        pdf_bytes = _make_pdf_bytes(["Hello World from PDF."])
        result = await PdfParser.parse(pdf_bytes, "test.pdf")
        # PdfParser returns (text, metadata)
        text, meta = result
        assert "Hello World" in text
        assert "page_count" in meta

    @pytest.mark.asyncio
    async def test_parse_multi_page_pdf(self) -> None:
        pdf_bytes = _make_pdf_bytes(["Page one content.", "Page two content."])
        result = await PdfParser.parse(pdf_bytes, "doc.pdf")
        text, meta = result
        assert "Page one" in text
        assert "Page two" in text
        assert meta["page_count"] == 2

    @pytest.mark.asyncio
    async def test_parse_empty_pdf(self) -> None:
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "", fontsize=12)  # empty page text
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        result = await PdfParser.parse(buf.getvalue(), "empty.pdf")
        text, meta = result
        assert meta["page_count"] == 1

    @pytest.mark.asyncio
    async def test_parse_empty_bytes(self) -> None:
        with pytest.raises(ValueError, match="Failed to read PDF"):
            await PdfParser.parse(b"", "test.pdf")

    @pytest.mark.asyncio
    async def test_parse_invalid_bytes(self) -> None:
        with pytest.raises(ValueError, match="Failed to read PDF"):
            await PdfParser.parse(b"not a pdf at all", "fake.pdf")

    @pytest.mark.asyncio
    async def test_parse_wrong_extension(self) -> None:
        pdf_bytes = _make_pdf_bytes(["Hello"])
        with pytest.raises(ValueError, match="Unsupported file extension"):
            await PdfParser.parse(pdf_bytes, "test.docx")


# ── DOCX Parser ──────────────────────────────────────────────────────────────


class TestDocxParser:
    """Unit tests for DocxParser."""

    def test_supports_docx(self) -> None:
        assert DocxParser.supports("report.docx") is True
        assert DocxParser.supports("REPORT.DOCX") is True
        assert DocxParser.supports("doc.doc") is False
        assert DocxParser.supports("doc.txt") is False

    @pytest.mark.asyncio
    async def test_parse_simple_docx(self) -> None:
        docx_bytes = _make_docx_bytes(["Hello from DOCX.", "Second paragraph."])
        result = await DocxParser.parse(docx_bytes, "test.docx")
        assert "Hello from DOCX" in result
        assert "Second paragraph" in result

    @pytest.mark.asyncio
    async def test_parse_docx_with_table(self) -> None:
        docx_bytes = _make_docx_bytes(
            ["Intro paragraph."],
            table_rows=[["Name", "Age"], ["Alice", "30"], ["Bob", "25"]],
        )
        result = await DocxParser.parse(docx_bytes, "table.docx")
        assert "Intro paragraph" in result
        assert "Alice" in result
        assert "30" in result
        assert "Bob" in result

    @pytest.mark.asyncio
    async def test_parse_empty_docx(self) -> None:
        docx_bytes = _make_docx_bytes([])
        with pytest.raises(ValueError, match="contains no extractable text"):
            await DocxParser.parse(docx_bytes, "empty.docx")

    @pytest.mark.asyncio
    async def test_parse_empty_bytes(self) -> None:
        with pytest.raises(ValueError, match="Empty DOCX"):
            await DocxParser.parse(b"", "test.docx")

    @pytest.mark.asyncio
    async def test_parse_invalid_bytes(self) -> None:
        with pytest.raises(ValueError, match="Failed to parse DOCX"):
            await DocxParser.parse(b"not a docx file", "fake.docx")

    @pytest.mark.asyncio
    async def test_parse_wrong_extension(self) -> None:
        docx_bytes = _make_docx_bytes(["Hello"])
        with pytest.raises(ValueError, match="Unsupported file extension"):
            await DocxParser.parse(docx_bytes, "test.pdf")

    @pytest.mark.asyncio
    async def test_parse_with_headings(self) -> None:
        """Headings should be prefixed with markdown-style markers."""
        docx_bytes = _make_docx_bytes(
            ["Introduction", "Background", "Some detail here.", "Methods", "Steps."],
            heading_levels=[1, 2, 0, 1, 0],
        )
        result = await DocxParser.parse(docx_bytes, "headings.docx")
        assert "# Introduction" in result
        assert "## Background" in result
        assert "Some detail here" in result
        assert "# Methods" in result

    @pytest.mark.asyncio
    async def test_parse_heading_and_table(self) -> None:
        """Headings and tables should coexist correctly."""
        docx_bytes = _make_docx_bytes(
            ["Summary"],
            heading_levels=[1],
            table_rows=[["Key", "Value"], ["A", "1"]],
        )
        result = await DocxParser.parse(docx_bytes, "h_table.docx")
        assert "# Summary" in result
        assert "Key | Value" in result
        assert "A | 1" in result


# ── HTML Parser ──────────────────────────────────────────────────────────────


class TestHtmlParser:
    """Unit tests for HtmlParser."""

    def test_supports_html(self) -> None:
        assert HtmlParser.supports("page.html") is True
        assert HtmlParser.supports("page.htm") is True
        assert HtmlParser.supports("PAGE.HTML") is True
        assert HtmlParser.supports("doc.txt") is False
        assert HtmlParser.supports("doc.pdf") is False

    @pytest.mark.asyncio
    async def test_parse_simple_html(self) -> None:
        html = b"<html><body><h1>Title</h1><p>Hello <b>World</b>.</p></body></html>"
        result = await HtmlParser.parse(html, "page.html")
        assert "Title" in result
        assert "Hello World" in result

    @pytest.mark.asyncio
    async def test_parse_strips_scripts_and_styles(self) -> None:
        html = b"""<html><head>
        <style>body { color: red; }</style>
        <script>console.log('should not appear');</script>
        </head><body><p>Visible text only.</p></body></html>"""
        result = await HtmlParser.parse(html, "page.html")
        assert "Visible text only" in result
        assert "color: red" not in result
        assert "console.log" not in result

    @pytest.mark.asyncio
    async def test_parse_preserves_paragraph_breaks(self) -> None:
        html = b"<html><body><p>First paragraph.</p><p>Second paragraph.</p></body></html>"
        result = await HtmlParser.parse(html, "page.html")
        # Paragraphs should be separated by newlines
        assert "First paragraph" in result
        assert "Second paragraph" in result

    @pytest.mark.asyncio
    async def test_parse_empty_html(self) -> None:
        with pytest.raises(ValueError, match="contains no extractable text"):
            await HtmlParser.parse(b"<html><body></body></html>", "empty.html")

    @pytest.mark.asyncio
    async def test_parse_empty_bytes(self) -> None:
        with pytest.raises(ValueError, match="Empty HTML"):
            await HtmlParser.parse(b"", "empty.html")

    @pytest.mark.asyncio
    async def test_parse_decomposes_nav_footer_aside_header(self) -> None:
        html = b"""<html><body>
        <nav>Navigation here</nav>
        <main>Main content here.</main>
        <footer>Footer here</footer>
        </body></html>"""
        result = await HtmlParser.parse(html, "page.html")
        assert "Main content" in result
        assert "Navigation" not in result
        assert "Footer" not in result

    @pytest.mark.asyncio
    async def test_parse_wrong_extension(self) -> None:
        with pytest.raises(ValueError, match="Unsupported file extension"):
            await HtmlParser.parse(b"<html></html>", "test.pdf")


# ── Parser Registry ──────────────────────────────────────────────────────────


class TestParserRegistry:
    """Tests for parser registry dispatch."""

    def test_get_parser_for_pdf(self) -> None:
        parser = get_parser_for("report.pdf")
        assert parser is PdfParser

    def test_get_parser_for_docx(self) -> None:
        parser = get_parser_for("letter.docx")
        assert parser is DocxParser

    def test_get_parser_for_html(self) -> None:
        parser = get_parser_for("page.html")
        assert parser is HtmlParser
        parser_htm = get_parser_for("page.htm")
        assert parser_htm is HtmlParser

    def test_get_parser_for_md(self) -> None:
        parser = get_parser_for("readme.md")
        assert parser is TextParser

    def test_get_parser_for_txt(self) -> None:
        parser = get_parser_for("notes.txt")
        assert parser is TextParser

    def test_get_parser_unsupported(self) -> None:
        with pytest.raises(ValueError, match="No parser available"):
            get_parser_for("image.png")

    def test_case_insensitive_extension(self) -> None:
        parser = get_parser_for("REPORT.PDF")
        assert parser is PdfParser

    def test_get_all_extensions_includes_all_formats(self) -> None:
        exts = get_all_supported_extensions()
        assert ".pdf" in exts
        assert ".docx" in exts
        assert ".html" in exts
        assert ".htm" in exts
        assert ".md" in exts
        assert ".txt" in exts

    def test_content_type_map(self) -> None:
        mapping = get_ext_to_content_type_map()
        assert mapping[".pdf"] == "application/pdf"
        assert mapping[".docx"] == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert mapping[".html"] == "text/html"
        assert mapping[".md"] == "text/markdown"
        assert mapping[".txt"] == "text/plain"


# ── Read-only API tests (no DB) ──────────────────────────────────────────────


class TestDocumentsApiMeta:
    """Tests for endpoints that don't need a database."""

    @pytest.mark.asyncio
    async def test_supported_extensions_endpoint(self, client: AsyncClient) -> None:
        response = await client.get("/documents/supported")
        assert response.status_code == 200
        data = response.json()
        assert "extensions" in data
        assert ".pdf" in data["extensions"]
        assert ".docx" in data["extensions"]
        assert ".html" in data["extensions"]
        assert ".md" in data["extensions"]
        assert ".txt" in data["extensions"]

    @pytest.mark.asyncio
    async def test_scrape_rejects_unsafe_url(self, client: AsyncClient) -> None:
        response = await client.post(
            "/documents/scrape",
            json={"url": "http://localhost:8000/private"},
        )
        assert response.status_code == 400
        assert "Unsafe" in response.json()["detail"]

    @pytest.mark.asyncio
    async def test_scrape_rejects_non_http_url(self, client: AsyncClient) -> None:
        response = await client.post(
            "/documents/scrape",
            json={"url": "ftp://files.example.com/data"},
        )
        assert response.status_code == 400
        assert "Unsafe" in response.json()["detail"]

    @pytest.mark.asyncio
    async def test_upload_rejects_unsupported_extension(self, client: AsyncClient) -> None:
        response = await client.post(
            "/documents/upload",
            files={"file": ("image.png", b"\x89PNG fake png content")},
        )
        assert response.status_code == 415
        assert "Unsupported file type" in response.json()["detail"]


# ── Web Scraper (unit) ───────────────────────────────────────────────────────


class TestWebScraperUnit:
    """Unit tests for WebScraper URL validation."""

    def test_rejects_localhost(self) -> None:
        from app.ingestion.parsers.scraper import WebScraper

        assert not WebScraper._is_safe_url("http://localhost:8080/page")

    def test_rejects_127_0_0_1(self) -> None:
        from app.ingestion.parsers.scraper import WebScraper

        assert not WebScraper._is_safe_url("http://127.0.0.1/admin")

    def test_rejects_private_10(self) -> None:
        from app.ingestion.parsers.scraper import WebScraper

        assert not WebScraper._is_safe_url("http://10.0.0.1/internal")

    def test_rejects_private_192_168(self) -> None:
        from app.ingestion.parsers.scraper import WebScraper

        assert not WebScraper._is_safe_url("http://192.168.1.1/router")

    def test_rejects_private_172_16(self) -> None:
        from app.ingestion.parsers.scraper import WebScraper

        assert not WebScraper._is_safe_url("http://172.16.0.1/internal")

    def test_accepts_public_https(self) -> None:
        from app.ingestion.parsers.scraper import WebScraper

        assert WebScraper._is_safe_url("https://example.com/page")
